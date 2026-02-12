import os
import logging
import shutil
from typing import Dict, Any, List, Optional
from datetime import datetime
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from utils.database import get_db
from modules.document.models import Document, DocumentVersion

logger = logging.getLogger(__name__)

class DocumentManager:
    """文档管理器"""
    
    def __init__(self):
        self.upload_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'uploads')
        os.makedirs(self.upload_dir, exist_ok=True)
    
    def upload_document(self, file_path: str, data: Dict[str, Any]) -> Document:
        """上传文档
        
        Args:
            file_path: 文件路径
            data: 文档数据
            
        Returns:
            上传的文档
        """
        db = next(get_db())
        try:
            # 生成文档ID
            document_id = f"DOC_{datetime.now().strftime('%Y%m%d%H%M%S')}_{os.urandom(4).hex()}"
            
            # 获取文件信息
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            file_type = self._get_file_type(file_name)
            
            # 验证文件类型
            if not self._is_allowed_file_type(file_type):
                raise ValueError(f"不支持的文件类型: {file_type}")
            
            # 验证文件大小
            max_file_size = 10 * 1024 * 1024  # 10MB
            if file_size > max_file_size:
                raise ValueError(f"文件大小超过限制: {file_size} > {max_file_size}")
            
            # 保存文件
            save_path = self._save_file(file_path, document_id, file_name)
            
            # 创建文档记录
            document = Document(
                document_id=document_id,
                customer_id=data.get('customer_id'),
                title=data.get('title', file_name),
                description=data.get('description'),
                file_name=file_name,
                file_path=save_path,
                file_size=file_size,
                file_type=file_type,
                created_by=data.get('created_by')
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # 创建初始版本
            self._create_initial_version(document)
            
            # 处理文档内容
            if self._needs_ocr(file_type):
                # 执行OCR
                self._process_ocr(document)
            else:
                # 提取文本内容
                self._extract_text(document)
            
            logger.info(f"上传文档成功: {file_name}")
            return document
            
        except Exception as e:
            db.rollback()
            logger.error(f"上传文档时出错: {e}")
            raise
        finally:
            db.close()
    
    def _save_file(self, file_path: str, document_id: str, file_name: str) -> str:
        """保存文件
        
        Args:
            file_path: 原文件路径
            document_id: 文档ID
            file_name: 文件名
            
        Returns:
            保存路径
        """
        # 创建文档目录
        doc_dir = os.path.join(self.upload_dir, document_id)
        os.makedirs(doc_dir, exist_ok=True)
        
        # 保存文件
        save_path = os.path.join(doc_dir, file_name)
        shutil.copy2(file_path, save_path)
        
        return save_path
    
    def _get_file_type(self, file_name: str) -> str:
        """获取文件类型
        
        Args:
            file_name: 文件名
            
        Returns:
            文件类型
        """
        ext = os.path.splitext(file_name)[1].lower()
        return ext[1:] if ext else 'unknown'
    
    def _is_allowed_file_type(self, file_type: str) -> bool:
        """检查是否是允许的文件类型
        
        Args:
            file_type: 文件类型
            
        Returns:
            是否允许
        """
        allowed_types = ['pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png']
        return file_type in allowed_types
    
    def _needs_ocr(self, file_type: str) -> bool:
        """检查是否需要OCR
        
        Args:
            file_type: 文件类型
            
        Returns:
            是否需要OCR
        """
        image_types = ['jpg', 'jpeg', 'png']
        return file_type in image_types
    
    def _create_initial_version(self, document: Document):
        """创建初始版本
        
        Args:
            document: 文档对象
        """
        db = next(get_db())
        try:
            version = DocumentVersion(
                document_id=document.id,
                version=1,
                file_name=document.file_name,
                file_path=document.file_path,
                file_size=document.file_size,
                created_by=document.created_by
            )
            db.add(version)
            db.commit()
        finally:
            db.close()
    
    def _process_ocr(self, document: Document):
        """执行OCR
        
        Args:
            document: 文档对象
        """
        try:
            # 打开图片
            image = Image.open(document.file_path)
            
            # 执行OCR
            text = pytesseract.image_to_string(image, lang='chi_sim')
            
            # 保存OCR结果
            ocr_path = os.path.join(os.path.dirname(document.file_path), 'ocr_result.txt')
            with open(ocr_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            # 更新文档状态
            db = next(get_db())
            try:
                doc = db.query(Document).filter(Document.id == document.id).first()
                doc.ocr_processed = True
                doc.content_extracted = True
                db.commit()
            finally:
                db.close()
            
            logger.info(f"OCR处理成功: {document.file_name}")
            
        except Exception as e:
            logger.error(f"OCR处理时出错: {e}")
    
    def _extract_text(self, document: Document):
        """提取文本内容
        
        Args:
            document: 文档对象
        """
        try:
            file_type = document.file_type
            file_path = document.file_path
            
            text = ""
            
            if file_type == 'pdf':
                text = self._extract_pdf_text(file_path)
            elif file_type in ['doc', 'docx']:
                text = self._extract_doc_text(file_path)
            
            # 保存提取结果
            if text:
                text_path = os.path.join(os.path.dirname(file_path), 'extracted_text.txt')
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                # 更新文档状态
                db = next(get_db())
                try:
                    doc = db.query(Document).filter(Document.id == document.id).first()
                    doc.content_extracted = True
                    db.commit()
                finally:
                    db.close()
                
                logger.info(f"文本提取成功: {document.file_name}")
                
        except Exception as e:
            logger.error(f"提取文本时出错: {e}")
    
    def _extract_pdf_text(self, pdf_path: str) -> str:
        """提取PDF文本
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            提取的文本
        """
        text = ""
        with open(pdf_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + '\n'
        return text
    
    def _extract_doc_text(self, doc_path: str) -> str:
        """提取Word文本
        
        Args:
            doc_path: Word文件路径
            
        Returns:
            提取的文本
        """
        text = ""
        if doc_path.endswith('.docx'):
            doc = DocxDocument(doc_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
        return text
    
    def update_document(self, document_id: int, data: Dict[str, Any]) -> Document:
        """更新文档
        
        Args:
            document_id: 文档ID
            data: 更新数据
            
        Returns:
            更新后的文档
        """
        db = next(get_db())
        try:
            # 查找文档
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise ValueError(f"文档不存在: {document_id}")
            
            # 更新字段
            for key, value in data.items():
                if hasattr(document, key):
                    setattr(document, key, value)
            
            document.updated_at = datetime.now()
            db.commit()
            db.refresh(document)
            
            logger.info(f"更新文档成功: {document.title}")
            return document
            
        except Exception as e:
            db.rollback()
            logger.error(f"更新文档时出错: {e}")
            raise
        finally:
            db.close()
    
    def delete_document(self, document_id: int) -> bool:
        """删除文档
        
        Args:
            document_id: 文档ID
            
        Returns:
            是否删除成功
        """
        db = next(get_db())
        try:
            # 查找文档
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise ValueError(f"文档不存在: {document_id}")
            
            # 删除文件
            if os.path.exists(document.file_path):
                # 删除整个文档目录
                doc_dir = os.path.dirname(document.file_path)
                if os.path.exists(doc_dir):
                    shutil.rmtree(doc_dir)
            
            # 删除版本
            db.query(DocumentVersion).filter(DocumentVersion.document_id == document_id).delete()
            
            # 删除文档
            db.delete(document)
            db.commit()
            
            logger.info(f"删除文档成功: {document.title}")
            return True
            
        except Exception as e:
            db.rollback()
            logger.error(f"删除文档时出错: {e}")
            raise
        finally:
            db.close()
    
    def get_document(self, document_id: int) -> Optional[Document]:
        """获取文档
        
        Args:
            document_id: 文档ID
            
        Returns:
            文档
        """
        db = next(get_db())
        try:
            return db.query(Document).filter(Document.id == document_id).first()
        finally:
            db.close()
    
    def list_documents(self, **filters) -> List[Document]:
        """列出文档
        
        Args:
            filters: 过滤条件
            
        Returns:
            文档列表
        """
        db = next(get_db())
        try:
            query = db.query(Document)
            
            # 应用过滤条件
            if 'customer_id' in filters:
                query = query.filter(Document.customer_id == filters['customer_id'])
            if 'status' in filters:
                query = query.filter(Document.status == filters['status'])
            if 'file_type' in filters:
                query = query.filter(Document.file_type == filters['file_type'])
            
            # 排序
            query = query.order_by(Document.created_at.desc())
            
            return query.all()
        finally:
            db.close()
    
    def get_document_versions(self, document_id: int) -> List[DocumentVersion]:
        """获取文档版本
        
        Args:
            document_id: 文档ID
            
        Returns:
            版本列表
        """
        db = next(get_db())
        try:
            return db.query(DocumentVersion).filter(
                DocumentVersion.document_id == document_id
            ).order_by(DocumentVersion.version.desc()).all()
        finally:
            db.close()
    
    def analyze_document(self, document_id: int) -> Dict[str, Any]:
        """分析文档
        
        Args:
            document_id: 文档ID
            
        Returns:
            分析结果
        """
        db = next(get_db())
        try:
            # 查找文档
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise ValueError(f"文档不存在: {document_id}")
            
            # 分析文档内容
            analysis_result = {
                'document_id': document.id,
                'title': document.title,
                'file_type': document.file_type,
                'file_size': document.file_size,
                'content_analysis': self._analyze_content(document),
                'risk_assessment': self._assess_risk(document)
            }
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"分析文档时出错: {e}")
            return {}
        finally:
            db.close()
    
    def _analyze_content(self, document: Document) -> Dict[str, Any]:
        """分析文档内容
        
        Args:
            document: 文档对象
            
        Returns:
            内容分析结果
        """
        # 读取提取的文本
        text_path = os.path.join(os.path.dirname(document.file_path), 'extracted_text.txt')
        ocr_path = os.path.join(os.path.dirname(document.file_path), 'ocr_result.txt')
        
        text = ""
        if os.path.exists(text_path):
            with open(text_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif os.path.exists(ocr_path):
            with open(ocr_path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        # 简单的内容分析
        analysis = {
            'text_length': len(text),
            'word_count': len(text.split()),
            'has_sensitive_info': self._detect_sensitive_info(text),
            'legal_terms': self._extract_legal_terms(text)
        }
        
        return analysis
    
    def _assess_risk(self, document: Document) -> Dict[str, Any]:
        """评估文档风险
        
        Args:
            document: 文档对象
            
        Returns:
            风险评估结果
        """
        # 读取提取的文本
        text_path = os.path.join(os.path.dirname(document.file_path), 'extracted_text.txt')
        ocr_path = os.path.join(os.path.dirname(document.file_path), 'ocr_result.txt')
        
        text = ""
        if os.path.exists(text_path):
            with open(text_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif os.path.exists(ocr_path):
            with open(ocr_path, 'r', encoding='utf-8') as f:
                text = f.read()
        
        # 简单的风险评估
        risk_level = 0
        risk_factors = []
        
        # 检测风险因素
        if '违约' in text:
            risk_level += 1
            risk_factors.append('违约条款')
        if '赔偿' in text:
            risk_level += 1
            risk_factors.append('赔偿条款')
        if '诉讼' in text:
            risk_level += 1
            risk_factors.append('诉讼风险')
        if '争议' in text:
            risk_level += 1
            risk_factors.append('争议解决')
        
        # 限制风险等级
        risk_level = min(risk_level, 3)
        
        return {
            'risk_level': risk_level,
            'risk_factors': risk_factors,
            'risk_description': self._get_risk_description(risk_level)
        }
    
    def _detect_sensitive_info(self, text: str) -> bool:
        """检测敏感信息
        
        Args:
            text: 文本
            
        Returns:
            是否包含敏感信息
        """
        sensitive_words = ['身份证', '银行卡', '密码', '隐私', '商业秘密']
        return any(word in text for word in sensitive_words)
    
    def _extract_legal_terms(self, text: str) -> List[str]:
        """提取法律条款
        
        Args:
            text: 文本
            
        Returns:
            法律条款列表
        """
        legal_terms = []
        
        # 简单的条款提取
        if '合同' in text:
            legal_terms.append('合同条款')
        if '协议' in text:
            legal_terms.append('协议条款')
        if '条款' in text:
            legal_terms.append('法律条款')
        
        return legal_terms
    
    def _get_risk_description(self, risk_level: int) -> str:
        """获取风险描述
        
        Args:
            risk_level: 风险等级
            
        Returns:
            风险描述
        """
        descriptions = [
            '无明显风险',
            '低风险',
            '中风险',
            '高风险'
        ]
        return descriptions[risk_level]

# 创建文档管理器实例
document_manager = DocumentManager()
